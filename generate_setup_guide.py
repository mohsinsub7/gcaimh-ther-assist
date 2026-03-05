"""Generate TherAssist Setup Guide Word document for Clinical Director."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ─── Page margins ───
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ─── Styles ───
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ─── Title Page ───
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('TherAssist')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Setup & Launch Guide')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('AI-Powered Therapy Session Assistant')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Version 2.0  |  March 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x80, 0x86, 0x8B)

doc.add_page_break()

# ─── Table of Contents ───
doc.add_heading('Table of Contents', level=1)
toc_items = [
    ('1.', 'What You Need Before Starting'),
    ('2.', 'Getting the TherAssist Folder'),
    ('3.', 'Launching TherAssist (Windows)'),
    ('4.', 'Launching TherAssist (Mac)'),
    ('5.', 'Using TherAssist'),
    ('6.', 'Stopping TherAssist'),
    ('7.', 'Launching Again (After First Time)'),
    ('8.', 'Troubleshooting'),
]
for num, item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}  {item}')
    run.font.size = Pt(12)

doc.add_page_break()

# ─── Section 1: What You Need ───
doc.add_heading('1. What You Need Before Starting', level=1)

p = doc.add_paragraph()
p.add_run('TherAssist runs on your computer and connects to SUNY\'s Google Cloud for AI analysis. '
           'Here is what you need:').font.size = Pt(11)

doc.add_paragraph()

items = [
    ('A computer', 'Windows 10/11 or Mac (macOS 12 or newer)'),
    ('Internet connection', 'Required for sign-in and AI analysis'),
    ('Your SUNY email', 'Your @downstate.edu email address and password'),
    ('The TherAssist folder', 'Provided by the development team (see next section)'),
]
for title_text, desc_text in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title_text)
    run.bold = True
    p.add_run(f' \u2014 {desc_text}')

p = doc.add_paragraph()
run = p.add_run('\nImportant: ')
run.bold = True
run.font.color.rgb = RGBColor(0xD9, 0x30, 0x25)
p.add_run('You do NOT need to install any software manually. '
           'The launcher script will automatically install everything needed on your computer the first time you run it.')

doc.add_page_break()

# ─── Section 2: Getting the Folder ───
doc.add_heading('2. Getting the TherAssist Folder', level=1)

steps = [
    'You will receive a folder called "gcaimh-ther-assist-DEV" (via email, USB drive, or shared link).',
    'Save this folder somewhere easy to find, like your Desktop or Documents folder.',
    'Open the folder. You should see files including:\n'
    '    \u2022  START-Windows.bat  (for Windows)\n'
    '    \u2022  START-Mac.command  (for Mac)\n'
    '    \u2022  gcloud-login-config.json\n'
    '    \u2022  Several subfolders (backend, frontend, etc.)',
    'Do NOT rename, move, or delete any files inside this folder.',
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {i}: ')
    run.bold = True
    p.add_run(step)

doc.add_page_break()

# ─── Section 3: Launching on Windows ───
doc.add_heading('3. Launching TherAssist (Windows)', level=1)

p = doc.add_paragraph()
p.add_run('Follow these steps to start TherAssist on a Windows computer.').font.size = Pt(11)

doc.add_heading('First-Time Launch (5\u201310 minutes)', level=2)

steps = [
    ('Open the TherAssist folder',
     'Navigate to the "gcaimh-ther-assist-DEV" folder on your computer.'),
    ('Double-click START-Windows.bat',
     'A black command window will appear with blue text. This is the launcher.'),
    ('Wait for software installation',
     'If this is your first time, the launcher will automatically install Python, Node.js, and '
     'Google Cloud SDK. This may take a few minutes. You will see progress messages.\n\n'
     'If it says "Software was installed successfully! Please CLOSE this window and double-click '
     'START-Windows.bat again" \u2014 do exactly that. Close the window and double-click the file again.'),
    ('Sign in with your SUNY account',
     'A sign-in link will appear in the command window.\n\n'
     '\u2022  Copy the link (right-click the command window to copy)\n'
     '\u2022  Paste it into any browser (Edge, Chrome, etc.)\n'
     '\u2022  Sign in with your @downstate.edu email and password\n'
     '\u2022  Complete any two-factor authentication if prompted\n'
     '\u2022  After signing in successfully, the command window will automatically continue'),
    ('Wait for services to start',
     'The launcher will set up the application services. On the first run, this takes 5\u201310 minutes '
     'because it needs to download and configure everything.\n\n'
     'You will see a progress display like:\n'
     '    [3/6] Setting up services...\n'
     '    therapy-analysis...         OK!\n'
     '    frontend...                 OK!\n\n'
     'Just wait \u2014 it will finish on its own.'),
    ('Browser opens automatically',
     'Once everything is ready, your default browser will open to http://localhost:3000 \u2014 '
     'this is TherAssist running on your computer.\n\n'
     'If the page appears blank, wait 30 seconds and press F5 to refresh.'),
]

for i, (title_text, desc_text) in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {i}: {title_text}')
    run.bold = True
    run.font.size = Pt(11)
    p = doc.add_paragraph(desc_text)

p = doc.add_paragraph()
run = p.add_run('\u26A0  Keep the black command window open!')
run.bold = True
run.font.color.rgb = RGBColor(0xD9, 0x30, 0x25)
p = doc.add_paragraph('The command window must stay open while you use TherAssist. '
                       'It is running the application in the background. '
                       'Closing it will stop TherAssist.')

doc.add_page_break()

# ─── Section 4: Launching on Mac ───
doc.add_heading('4. Launching TherAssist (Mac)', level=1)

p = doc.add_paragraph()
p.add_run('Follow these steps to start TherAssist on a Mac.').font.size = Pt(11)

doc.add_heading('One-Time Setup: Allow the Script to Run', level=2)

p = doc.add_paragraph('Before the first launch, you need to allow the script to run:')

steps_mac_setup = [
    'Open the TherAssist folder in Finder.',
    'Open Terminal (press Cmd+Space, type "Terminal", press Enter).',
    'In Terminal, type:  chmod +x \n'
    'Then drag the START-Mac.command file from Finder into the Terminal window.\n'
    'It will paste the file path automatically. Press Enter.',
    'Close Terminal. Now go back to Finder and double-click START-Mac.command.',
    'If a security warning appears saying the app is from an unidentified developer, '
    'click "Open" to confirm. This only needs to be done once.',
]
for i, step in enumerate(steps_mac_setup, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    p.add_run(step)

doc.add_heading('First-Time Launch (5\u201310 minutes)', level=2)

steps_mac = [
    ('Double-click START-Mac.command',
     'A Terminal window will appear. This is the launcher.'),
    ('Allow Homebrew installation (if needed)',
     'If this is your first time, the launcher will install Homebrew (a Mac package manager), '
     'then Python, Node.js, and Google Cloud SDK.\n\n'
     'You may be asked for your Mac password (the one you use to log in to your Mac). '
     'Type it and press Enter. The characters will not appear as you type \u2014 this is normal.\n\n'
     'If it says "Software was installed successfully! Please CLOSE this window..." \u2014 '
     'close the Terminal and double-click START-Mac.command again.'),
    ('Sign in with your SUNY account',
     'A sign-in link will appear in the Terminal window.\n\n'
     '\u2022  Copy the link (Cmd+C or right-click \u2192 Copy)\n'
     '\u2022  Paste it into any browser (Safari, Chrome, etc.)\n'
     '\u2022  Sign in with your @downstate.edu email and password\n'
     '\u2022  Complete any two-factor authentication if prompted\n'
     '\u2022  After signing in successfully, the Terminal will automatically continue'),
    ('Wait for services to start',
     'Same as Windows \u2014 the first run takes 5\u201310 minutes to download and set up everything. '
     'Just wait for it to finish.'),
    ('Browser opens automatically',
     'Your default browser will open to http://localhost:3000 with TherAssist ready to use.'),
]

for i, (title_text, desc_text) in enumerate(steps_mac, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {i}: {title_text}')
    run.bold = True
    run.font.size = Pt(11)
    p = doc.add_paragraph(desc_text)

p = doc.add_paragraph()
run = p.add_run('\u26A0  Keep the Terminal window open!')
run.bold = True
run.font.color.rgb = RGBColor(0xD9, 0x30, 0x25)
p = doc.add_paragraph('The Terminal window must stay open while you use TherAssist.')

doc.add_page_break()

# ─── Section 5: Using TherAssist ───
doc.add_heading('5. Using TherAssist', level=1)

p = doc.add_paragraph('Once the browser opens to TherAssist:')

steps_use = [
    'You will see the TherAssist interface in your browser.',
    'Click "New Session" to start a therapy session.',
    'Select the therapy modality (CBT, DBT, IPT, or BA) and enter patient details.',
    'Use the microphone button to record live, or upload a pre-recorded audio file (.wav, .mp3, .m4a).',
    'TherAssist will transcribe speech in real-time and display it on the left panel.',
    'The AI provides real-time guidance (right panel): clinical metrics, suggested interventions, '
    'safety alerts, and pathway recommendations \u2014 all grounded in evidence-based therapy literature.',
    'The "LLM Activity" log at the bottom shows every AI analysis job (click to collapse/expand).',
    'When done, click "End Session" to generate a comprehensive session summary. '
    'You can save the session for later review.',
]
for i, step in enumerate(steps_use, 1):
    p = doc.add_paragraph(f'{i}. {step}')

p = doc.add_paragraph()
run = p.add_run('Note: ')
run.bold = True
p.add_run('Your browser must allow microphone access for the transcription to work. '
           'If prompted, click "Allow".')

doc.add_page_break()

# ─── Section 6: Stopping TherAssist ───
doc.add_heading('6. Stopping TherAssist', level=1)

doc.add_heading('Windows:', level=2)
steps_stop_win = [
    'Go to the black command window (it should still be open).',
    'Press any key on your keyboard.',
    'The window will show "All services stopped" and "Thank you for using TherAssist!"',
    'Press any key again to close the window.',
]
for i, step in enumerate(steps_stop_win, 1):
    p = doc.add_paragraph(f'{i}. {step}')

doc.add_heading('Mac:', level=2)
steps_stop_mac = [
    'Go to the Terminal window (it should still be open).',
    'Press Enter (Return) on your keyboard.',
    'The window will show "All services stopped" and "Thank you for using TherAssist!"',
    'Close the Terminal window.',
]
for i, step in enumerate(steps_stop_mac, 1):
    p = doc.add_paragraph(f'{i}. {step}')

doc.add_page_break()

# ─── Section 7: Launching Again ───
doc.add_heading('7. Launching Again (After First Time)', level=1)

p = doc.add_paragraph('After the first launch, starting TherAssist is much faster (under 1 minute):')

steps_again = [
    'Double-click START-Windows.bat (or START-Mac.command on Mac).',
    'Copy the sign-in link from the window and paste it into any browser. Sign in with your SUNY credentials.',
    'Wait about 30\u201360 seconds for services to start.',
    'The browser opens automatically \u2014 you\'re ready to go!',
]
for i, step in enumerate(steps_again, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    p.add_run(step)

p = doc.add_paragraph()
run = p.add_run('Why is it faster? ')
run.bold = True
p.add_run('The first launch downloads and installs all the software and dependencies. '
           'After that, everything is already on your computer, so it just needs to sign in and start up.')

doc.add_page_break()

# ─── Section 8: Troubleshooting ───
doc.add_heading('8. Troubleshooting', level=1)

p = doc.add_paragraph('If something goes wrong, the launcher will show an error message with:')
items = [
    'A description of what happened',
    'Step-by-step instructions to fix it',
    'An error code (like "AUTH-FAILED" or "PIP-ANALYSIS-FAILED")',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('\nIf you need help, send the following to Mohsin:')
run.bold = True

items_send = [
    'The error code shown on screen',
    'The file "error-log.txt" from the TherAssist folder (it\'s created automatically)',
    'A screenshot of the error (if possible)',
]
for item in items_send:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
doc.add_heading('Common Issues and Solutions', level=2)

# Table of common issues
table = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
table.columns[0].width = Inches(2.0)
table.columns[1].width = Inches(2.5)
table.columns[2].width = Inches(2.0)

hdr = table.rows[0].cells
hdr[0].text = 'What You See'
hdr[1].text = 'What It Means'
hdr[2].text = 'What To Do'
for cell in hdr:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

issues = [
    ('"Software was installed successfully!\nPlease CLOSE this window..."',
     'New programs were installed and your computer needs to recognize them.',
     'Close the window and double-click the START file again. This only happens once.'),
    ('"SIGN-IN FAILED"',
     'The Google sign-in didn\'t complete.',
     'Close the window, double-click START again, and make sure to sign in with your @downstate.edu email.'),
    ('"MISSING-LOGIN-CONFIG"',
     'A required file is missing from the folder.',
     'Contact Mohsin to get a fresh copy of the TherAssist folder.'),
    ('Page is blank after browser opens',
     'Services are still starting up.',
     'Wait 30 seconds and press F5 (Windows) or Cmd+R (Mac) to refresh.'),
    ('"INSTALL-PYTHON-FAILED"\nor similar install error',
     'Automatic installation didn\'t work.',
     'Follow the manual install instructions shown on screen, then try again.'),
    ('Services start but app doesn\'t work',
     'Your SUNY credentials may have expired.',
     'Close everything, double-click START again. The sign-in will refresh your credentials.'),
    ('"SLOW-START"',
     'Services took longer than 90 seconds.',
     'On first run this is normal. If the browser opened, just wait and refresh. '
     'If it keeps happening, try closing and restarting.'),
]

for what, means, fix in issues:
    row = table.add_row().cells
    row[0].text = what
    row[1].text = means
    row[2].text = fix

# ─── Contact ───
doc.add_paragraph()
doc.add_heading('Need More Help?', level=2)
p = doc.add_paragraph('Contact Mohsin Sardar:')
doc.add_paragraph('Email: mohsin.sardar@downstate.edu', style='List Bullet')

# ─── Save ───
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           'TherAssist_Setup_Guide_v2.docx')
doc.save(output_path)
print(f"Guide saved to: {output_path}")
